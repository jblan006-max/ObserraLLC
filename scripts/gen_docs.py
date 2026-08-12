#!/usr/bin/env python3
"""Generate the Obserra EU CRA Governance guides as PDF and Word (.docx).

Produces three role-targeted guides from ONE set of sections + screenshots:
  - Install & User Guide  (full)      -> Obserra-EU-CRA-Governance-Install-and-User-Guide.{pdf,docx}
  - Executive Guide       (short)     -> Obserra-EU-CRA-Governance-Executive-Guide.{pdf,docx}
  - Admin & Operator Guide(deep)      -> Obserra-EU-CRA-Governance-Admin-Operator-Guide.{pdf,docx}

Every guide opens with a branded cover (logo) + a numbered Contents page. Screenshots
are read from /app/scripts/shots and embedded.
"""
import os
from datetime import datetime

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(BASE, "scripts", "shots")
PUBLIC = os.path.join(BASE, "frontend", "public")
OUT = os.path.join(BASE, "backend", "assets", "docs")
os.makedirs(OUT, exist_ok=True)

NAVY = "#0f1e3d"
AI = "#12b4d6"
BRAND = "Obserra EU CRA Governance"
TAGLINE = "European Union Cyber Resilience Act (Regulation (EU) 2024/2847) product governance"

E, A = "exec", "admin"      # audience tags
ALL = (E, A)

# (heading, [paragraphs], screenshot_or_None, audiences)
SECTIONS = [
    ("About Obserra EU CRA Governance", [
        "Obserra EU CRA Governance helps manufacturers, importers and distributors bring products "
        "with digital elements into conformity with the European Union Cyber Resilience Act "
        "(Regulation (EU) 2024/2847). It gives compliance, security and executive teams a single live "
        "view of product classification, regulation-mapped readiness, software bills of materials "
        "(SBOMs), Article 14 vulnerability and incident reporting clocks, notified-body sign-off and "
        "the EU Declaration of Conformity that underpins CE marking.",
        "Every fact is computed LIVE from the underlying product records (No-Mock): proposed Class I / "
        "Class II / Critical / Default classification, the conformity assessment pathway, assessment "
        "scores and market-readiness gates are all derived from real records. Every regulatory action "
        "is written to a tamper-evident, hash-chained Internal Regulatory Ledger so the compliance "
        "timeline is audit-defensible, and secure vendor / notified-body portals never expose that "
        "private ledger.",
    ], None, ALL),

    ("Installing the App (One-Click PWA)", [
        "Obserra installs like a native app straight from the browser — no app store required — "
        "and works across desktop, tablet and mobile.",
        "Desktop (Chrome / Edge): click the Install icon in the address bar, or use the in-app "
        "'Install' banner. Android (Chrome): tap the 'Install' banner or menu -> Add to Home "
        "screen. iPhone / iPad (Safari): tap Share -> Add to Home Screen.",
        "Once installed, the app launches full-screen and can receive push notifications for "
        "Article 14 reporting deadlines, classification approvals and CE-readiness changes.",
    ], None, ALL),

    ("On-Premise Installation (Docker)", [
        "For fully self-hosted deployments, download the on-premise package from "
        "Settings -> Deployment & Documentation. An optional install.sh is included.",
        "Prerequisites: Docker 24+ and Docker Compose v2, 2 vCPU / 4 GB RAM (8 GB recommended).",
        "Steps: (1) place backend/ and frontend/ next to the deploy/ folder; (2) copy "
        ".env.example to .env and set JWT_SECRET, EMERGENT_LLM_KEY and PUBLIC_URL; (3) run "
        "docker compose -f deploy/docker-compose.yml --env-file deploy/.env up -d --build; "
        "(4) open http://<machine-ip>:8080. Full details are in the bundled INSTALL.md.",
    ], None, (A,)),

    ("Signing In", [
        "Open the app and sign in with your work email and password. Passwords follow "
        "NIST 800-63B (>=12 chars with upper/lower/number/symbol). Roles determine what you see: "
        "admins get full governance and approval controls, executives get the board-ready view.",
    ], "01_login.jpg", ALL),

    ("Getting Started — Load Sample Products", [
        "New tenants can populate the workspace instantly. On Products & Classification, an admin "
        "clicks 'Load Samples' to seed a small set of real, editable CRA product records — an "
        "identity broker, a web firewall, a smart lock, a secure element and a general productivity "
        "app — spanning Default, Class I, Class II and Critical classifications.",
        "The sample set also seeds an initial readiness assessment and a CycloneDX SBOM so the "
        "dashboard, AI Analyst and ledger tell a complete story on first open. Samples are clearly "
        "flagged and can be removed at any time with 'Clear Samples'; loading is idempotent, so it "
        "never creates duplicates.",
    ], None, ALL),

    ("Mission Control & the CRA AI Analyst", [
        "Mission Control is the board-ready rollup: products under governance, average "
        "regulation-mapped readiness, the Class I / II / Critical split, open external assessments, "
        "overdue Article 14 clocks and CE-ready count. Every tile is clickable and drills into the "
        "relevant workspace tab.",
        "The CRA AI Analyst sits at the top of Mission Control. It reads the LIVE CRA posture — "
        "product counts, classification split, named CE-marking blockers and overdue Article 14 "
        "reporting stages — and returns a concise executive briefing: a headline that opens with the "
        "countdown to the nearest statutory CRA deadline, three-to-five grounded insights labelled "
        "FACT / ESTIMATE / RISK, and prioritised recommended actions. Every statement is grounded in "
        "your own records; a deterministic fallback keeps the briefing available even if the AI "
        "service is momentarily unreachable. Click 'Regenerate' to refresh after making changes, and "
        "opt in to the weekly CRA AI Analyst email so leadership sees CE-readiness drift without "
        "logging in.",
    ], "cra_mission.jpg", ALL),

    ("Products & Classification", [
        "Register each product with digital elements: name, version, manufacturer legal name, "
        "description, core functionality and any known Annex III / IV category codes. Obserra then "
        "runs a deterministic, explainable classification: explicit Annex categories take priority, "
        "and where none are selected a transparent heuristic proposes Default, Class I, Class II or "
        "Critical with the matching conformity pathway (self-assessment, third-party assessment, or "
        "critical-product route).",
        "Classification is always PROPOSED until an authorised user records a formal approval with a "
        "written rationale. Approving locks the classification into the compliance record and writes "
        "the decision, actor and legal basis to the Internal Regulatory Ledger. Products can be "
        "re-classified whenever functionality or categories change.",
    ], "cra_products.jpg", ALL),

    ("CRA Readiness Assessments & the Certification Portal", [
        "Each product is assessed against a regulation-mapped requirement set. Every question links "
        "back to specific CRA Articles and Annex I essential requirements, and the score reflects "
        "conforming, partial and non-conforming answers with evidence references.",
        "Assessments can be completed internally or delegated. From the Certification Portal, an "
        "admin issues a tenant-scoped, time-limited link to a vendor (for self-assessment) or to a "
        "testing lab / notified body (for external sign-off). Tokens are stored only as hashes, and "
        "the external portal exposes only the invited product context — never the private Internal "
        "Regulatory Ledger.",
    ], None, (A,)),

    ("SBOM & Software Components", [
        "Obserra generates a machine-readable Software Bill of Materials in CycloneDX 1.6 or "
        "SPDX 2.3 from supported dependency manifests (requirements.txt, package.json / "
        "package-lock.json, pom.xml). Each generated artifact records the component count and full "
        "component inventory.",
        "SBOM generation is logged to the Internal Regulatory Ledger and mapped to Annex I Part II(1), "
        "supporting the CRA obligation to identify and document the components contained in a product "
        "with digital elements.",
    ], None, (A,)),

    ("Article 14 Vulnerability & Incident Reporting", [
        "The reporting workspace tracks the statutory clocks for actively exploited vulnerabilities "
        "and severe incidents: the 24-hour early warning, the 72-hour notification and the final "
        "report. Deadlines are computed from the moment of awareness, and overdue stages are flagged "
        "in red across Mission Control and the AI Analyst.",
        "Article 14 reporting obligations apply from 11 September 2026. Obserra prepares and tracks "
        "the submission package and records each stage against the single reporting platform; it "
        "does not claim a regulatory submission has been made until an official submission or receipt "
        "reference is recorded, keeping the timeline honest and audit-defensible.",
    ], "cra_vuln.jpg", ALL),

    ("Labs & Notified Bodies (Conformity Assessment)", [
        "Maintain a provider-neutral registry of testing labs, CRA notified bodies and certification "
        "bodies, capturing NANDO identity, country and verification evidence. Class II and Critical "
        "products — and Class I products without a fully applied harmonised route — require "
        "third-party conformity assessment.",
        "Raise external conformity assessment requests against a product and provider, supporting "
        "Module B + Module C, Module H, EU cybersecurity certification and testing-evidence "
        "workflows. External assessors record their decision and findings through the secure portal, "
        "and the outcome flows back into the product's market-readiness gates.",
    ], None, (A,)),

    ("EU Declaration of Conformity & CE Marking", [
        "The Declaration & CE workspace evaluates market readiness against live records: an approved "
        "classification, a complete readiness assessment, a generated SBOM, any required notified-body "
        "sign-off and an approved EU Declaration of Conformity. Each open gate is shown as an explicit "
        "blocker or warning.",
        "A completed assessment does NOT equal a Declaration approved, CE ready, or product placed on "
        "the market. Approving the EU Declaration of Conformity requires an authorised signatory name "
        "and title; the approval, signatory and declaration reference are written to the Internal "
        "Regulatory Ledger before a product is reported as CE ready. General CRA application begins "
        "11 December 2027.",
    ], "cra_declaration.jpg", ALL),

    ("Internal Regulatory Ledger & Auditor Verification", [
        "Every regulatory action — product registration, classification proposal and approval, "
        "assessment, SBOM generation, reporting submission and declaration approval — is appended to "
        "a private, hash-chained Internal Regulatory Ledger. Each record carries the prior-record "
        "hash and its own hash, so any tampering is detectable.",
        "For independent assurance, an admin issues a read-only Auditor Verification link for a "
        "product from the Regulatory Ledger tab and shares it with one tap. The auditor or notified "
        "body opens the link and sees the product's classification, conformity route, CE status, the "
        "tamper-evident compliance timeline and a live re-computation of hash-chain integrity — "
        "without ever exposing the private ledger payloads. Links are time-limited and tenant-scoped.",
    ], "cra_ledger.jpg", ALL),

    ("Regulation Map", [
        "The Regulation Map is the authoritative requirement catalogue: each governance object links "
        "back to Regulation (EU) 2024/2847 and Commission Implementing Regulation (EU) 2025/2392, "
        "with the obligation, legal basis and expected evidence types. It is the traceability layer "
        "that makes every workflow defensible against the source law.",
    ], "cra_regulation.jpg", (A,)),

    ("Executive Brief & Board Reporting", [
        "One click generates the EU CRA Governance Executive Brief as a branded PDF — product "
        "posture, classification split, readiness, Article 14 status and CE readiness — suitable for "
        "leadership and board packs. It is compiled from live records with an explicit note that "
        "Obserra provides regulatory workflow and traceability, not legal advice.",
    ], None, (A,)),

    ("Settings, Branding & Deployment", [
        "Personal preferences (digest cadence, replay the guided tour) plus admin controls: alert "
        "recipients, custom branding (company name, logo, accent colour), a 'send me a test now' "
        "button, and the Deployment & Documentation downloads — the on-premise package and these "
        "guides in PDF and Word.",
    ], "cra_settings.jpg", (A,)),

    ("Support & Legal Boundary", [
        "For assistance contact your Obserra administrator. Classifications are automatically "
        "PROPOSED and require authorised approval; readiness scores and AI Analyst briefings are "
        "decision-support outputs and do not constitute legal advice or a guarantee of CRA "
        "conformity. Placing a product on the EU market remains the manufacturer's responsibility.",
    ], None, ALL),
]


def _img_size(path):
    try:
        from PIL import Image
        with Image.open(path) as im:
            return im.size
    except Exception:
        return (890, 267)


def _logo_path():
    for name in ("brand-lockup.png", "brand-wordmark.png", "brand-mark.png"):
        p = os.path.join(PUBLIC, name)
        if os.path.exists(p):
            return p
    return None


def build_pdf(path, sections, guide_title):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    HRFlowable, PageBreak)
    styles = getSampleStyleSheet()
    ctitle = ParagraphStyle("ctitle", parent=styles["Title"], fontSize=30, leading=34,
                            textColor=colors.HexColor(NAVY), alignment=1)
    cbrand = ParagraphStyle("cbrand", parent=styles["Title"], fontSize=13, leading=16,
                            textColor=colors.HexColor(AI), alignment=1, spaceAfter=2)
    csub = ParagraphStyle("csub", parent=styles["Normal"], fontSize=11.5, leading=15,
                          textColor=colors.HexColor(NAVY), alignment=1)
    cmeta = ParagraphStyle("cmeta", parent=styles["Normal"], fontSize=9, leading=13,
                           textColor=colors.grey, alignment=1)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=15, leading=19,
                        textColor=colors.HexColor(AI), spaceBefore=16, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=6)
    toc = ParagraphStyle("toc", parent=body, fontSize=11, leading=20, textColor=colors.HexColor(NAVY))
    cap = ParagraphStyle("cap", parent=body, fontSize=8, textColor=colors.grey)

    doc = SimpleDocTemplate(path, pagesize=LETTER, topMargin=0.85 * inch, bottomMargin=0.8 * inch,
                            title=f"{BRAND} — {guide_title}", author="Obserra")
    story = []
    # ---- Cover ----
    story.append(Spacer(1, 130))
    lp = _logo_path()
    if lp:
        iw, ih = _img_size(lp)
        w = 3.6 * inch
        img = RLImage(lp, width=w, height=w * ih / iw); img.hAlign = "CENTER"
        story += [img, Spacer(1, 34)]
    story += [Paragraph(BRAND, cbrand), Paragraph(guide_title, ctitle), Spacer(1, 12),
              Paragraph(TAGLINE, csub), Spacer(1, 26),
              HRFlowable(width="38%", color=colors.HexColor(AI), hAlign="CENTER"), Spacer(1, 16),
              Paragraph(datetime.now().strftime("%B %Y"), cmeta),
              Paragraph("Confidential — prepared for internal use.", cmeta), PageBreak()]
    # ---- Contents ----
    story += [Paragraph("Contents", h1), Spacer(1, 4)]
    for i, (head, _p, _s, _a) in enumerate(sections, 1):
        story.append(Paragraph(f"{i}.&nbsp;&nbsp;{head}", toc))
    story.append(PageBreak())
    # ---- Body ----
    for i, (head, paras, shot, _a) in enumerate(sections, 1):
        story.append(Paragraph(f"{i}. {head}", h1))
        for p in paras:
            story.append(Paragraph(p, body))
        if shot:
            fp = os.path.join(SHOTS, shot)
            if os.path.exists(fp):
                iw = 6.4 * inch
                img = RLImage(fp, width=iw, height=iw * 900.0 / 1440.0)
                story += [Spacer(1, 4), img, Paragraph(f"Figure: {head}", cap), Spacer(1, 6)]

    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7); canvas.setFillColor(colors.grey)
        canvas.drawCentredString(LETTER[0] / 2, 0.5 * inch,
                                 f"{BRAND} — {guide_title}  ·  Executive Protection & Intelligence LLC  ·  Confidential")
        if _doc.page > 1:
            canvas.drawRightString(LETTER[0] - 0.7 * inch, 0.5 * inch, str(_doc.page))
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_docx(path, sections, guide_title):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = Document()
    # ---- Cover ----
    lp = _logo_path()
    if lp:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(lp, width=Inches(3.6))
    b = d.add_paragraph(); b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = b.add_run(BRAND); rb.bold = True; rb.font.size = Pt(13); rb.font.color.rgb = RGBColor(0x12, 0xb4, 0xd6)
    t = d.add_heading(guide_title, level=0); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph(TAGLINE); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = d.add_paragraph(f"{datetime.now().strftime('%B %Y')}  ·  Confidential — prepared for internal use.")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in meta.runs:
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    d.add_page_break()
    # ---- Contents ----
    d.add_heading("Contents", level=1)
    for i, (head, _p, _s, _a) in enumerate(sections, 1):
        d.add_paragraph(f"{i}.  {head}")
    d.add_page_break()
    # ---- Body ----
    for i, (head, paras, shot, _a) in enumerate(sections, 1):
        hd = d.add_heading(f"{i}. {head}", level=1)
        for run in hd.runs:
            run.font.color.rgb = RGBColor(0x0f, 0x1e, 0x3d)
        for p in paras:
            d.add_paragraph(p)
        if shot:
            fp = os.path.join(SHOTS, shot)
            if os.path.exists(fp):
                d.add_picture(fp, width=Inches(6.2))
                capp = d.add_paragraph(f"Figure: {head}"); capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in capp.runs:
                    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    d.save(path)


def _build(sections, title, pdf_name, docx_name):
    pdf = os.path.join(OUT, pdf_name); docx = os.path.join(OUT, docx_name)
    build_pdf(pdf, sections, title)
    build_docx(docx, sections, title)
    return pdf, docx


def generate_all():
    """Regenerate all role guides. Keeps pdf/docx/pdf_size/docx_size = the full guide."""
    exec_secs = [s for s in SECTIONS if E in s[3]]
    full_pdf, full_docx = _build(SECTIONS, "Install & User Guide",
                                 "Obserra-EU-CRA-Governance-Install-and-User-Guide.pdf", "Obserra-EU-CRA-Governance-Install-and-User-Guide.docx")
    exec_pdf, exec_docx = _build(exec_secs, "Executive Guide",
                                 "Obserra-EU-CRA-Governance-Executive-Guide.pdf", "Obserra-EU-CRA-Governance-Executive-Guide.docx")
    admin_pdf, admin_docx = _build(SECTIONS, "Admin & Operator Guide",
                                   "Obserra-EU-CRA-Governance-Admin-Operator-Guide.pdf", "Obserra-EU-CRA-Governance-Admin-Operator-Guide.docx")
    return {"pdf": full_pdf, "docx": full_docx,
            "pdf_size": os.path.getsize(full_pdf), "docx_size": os.path.getsize(full_docx),
            "exec_pdf": exec_pdf, "exec_docx": exec_docx,
            "admin_pdf": admin_pdf, "admin_docx": admin_docx}


if __name__ == "__main__":
    r = generate_all()
    for k in ("pdf", "exec_pdf", "admin_pdf", "docx", "exec_docx", "admin_docx"):
        print(k, os.path.getsize(r[k]), r[k])
